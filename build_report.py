# -*- coding: utf-8 -*-
"""產生 VPP-RTS 作業完整報告（含報告講稿）為 docx。

執行：  .venv/bin/python build_report.py
輸出：  報告_VPP-RTS.docx
所有數據取自 output/*.json 與 input/*.json（即實際執行結果）。
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = "Microsoft JhengHei"   # 微軟正黑體；缺字時 Word 會自動以系統中文字型替代
MONO = "Consolas"

doc = Document()


# --------------------------------------------------------------------------- 字型 / 樣式
def _set_eastasia(style, font_name=CJK, size=None):
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        style.font.size = Pt(size)


_set_eastasia(doc.styles["Normal"], CJK, 10.5)
for sname in ("Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet",
              "List Number"):
    try:
        _set_eastasia(doc.styles[sname], CJK)
    except KeyError:
        pass

# 行距與段距微調
normal = doc.styles["Normal"]
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.18


def _run_cjk(run, font_name=CJK):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


# --------------------------------------------------------------------------- 內容輔助函式
def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        _run_cjk(r)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        _run_cjk(r)
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        _run_cjk(r)
    return p


def para(text="", bold=False, italic=False, size=10.5, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    _run_cjk(run)
    return p


def rich(parts, size=10.5, align=None):
    """parts: list of (text, bold) 或 (text, bold, color)。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for item in parts:
        text, bold = item[0], item[1]
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if len(item) > 2 and item[2] is not None:
            run.font.color.rgb = item[2]
        _run_cjk(run)
    return p


def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * level)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        _run_cjk(r)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    _run_cjk(r)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    _run_cjk(r)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    # 淺灰底
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(9)
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:ascii"), MONO)
    rpr.get_or_add_rFonts().set(qn("w:hAnsi"), MONO)
    return p


def _shade_cell(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def table(headers, rows, widths=None, header_fill="305496", font=9.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        pp = hdr[i].paragraphs[0]
        r = pp.add_run(htext)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _run_cjk(r)
        _shade_cell(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            pp = cells[i].paragraphs[0]
            r = pp.add_run(str(val))
            r.font.size = Pt(font)
            _run_cjk(r)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def spacer(pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


def pagebreak():
    doc.add_page_break()


BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC0, 0x39, 0x2B)


# =========================================================================== 封面
title = doc.add_heading("", level=0)
trun = title.add_run("虛擬電廠即時排程系統 (VPP-RTS)")
trun.font.size = Pt(24)
_run_cjk(trun)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("即時系統 (Real-Time Systems) 作業　完整實作報告 + 報告講稿")
sr.font.size = Pt(14)
sr.bold = True
sr.font.color.rgb = BLUE
_run_cjk(sr)

spacer(10)
para("國立成功大學　即時系統課程", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para("Level 1（日前排程 + 接收測試 + 評估）＋ Level 2（放寬假設 + 進階動態排程）",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
spacer(20)

para("本文件用途", bold=True, size=12, color=BLUE)
para("本報告分為兩大部分：")
bullet("第一部分「完整實作說明」──讓未參與開發的組員，讀完即可掌握整份作業的問題定義、"
       "數學模型、四個階段的演算法、Level 2 的放寬假設與動態排程、以及實際執行結果。")
bullet("第二部分「報告講稿」──依投影片順序撰寫，可直接照著對老師口頭報告（含開場、各段重點、"
       "預期問答）。")
spacer(6)
para("註：報告中所有數據均取自專案 output/ 與 input/ 之實際 JSON 檔（單一隨機種子、可重現）。",
     italic=True, size=9.5)

pagebreak()

# =========================================================================== 目錄式概覽
h1("文件導覽")
para("第一部分　完整實作說明", bold=True, color=BLUE)
bullet("1. 作業背景與目標")
bullet("2. 系統輸入與資產設定")
bullet("3. 系統架構與管線（Pipeline）")
bullet("4. Phase 1：週期任務集生成")
bullet("5. Phase 2：日前 MILP 排程（決策變數／目標函數／23 條限制式）")
bullet("6. Phase 3：接收測試（Sporadic／Aperiodic）")
bullet("7. Phase 4：效能評估指標")
bullet("8. Level 2-A：放寬假設建模（R1–R10）")
bullet("9. Level 2-B：進階動態排程（滾動視窗再最佳化）")
bullet("10. 實際執行結果與靜態 vs 動態比較")
bullet("11. 程式碼結構與執行方式")
bullet("12. 對應評分表的自我檢核")
spacer(4)
para("第二部分　報告講稿（明日對老師報告）", bold=True, color=BLUE)
bullet("依 11 段投影片組織的逐字／重點講稿，附時間配置與預期問答")

pagebreak()

# =========================================================================== 第一部分
h1("第一部分　完整實作說明")

# ---- 1. 背景 ----
h2("1. 作業背景與目標")
para("「虛擬電廠（Virtual Power Plant, VPP）」把分散的發電與儲能資產整合成一個可被統一調度的"
     "電力系統。本作業要為一個 VPP 設計一套即時排程系統：它一方面要安排各種「發電設備」"
     "供電，另一方面要滿足一連串有時間限制的「用電需求（job）」。")
rich([("我們把發電設備視為即時系統中的「處理器（processor）」，把用電需求視為「任務（task/job）」，"
       "於是電力調度問題就被轉化成一個即時排程問題。", False)])
spacer(2)
para("三類發電設備（處理器）：", bold=True)
bullet("傳統機組 Ig：可控輸出的火力機組，有開關機、最小開／關機時間、爬升率與成本。", bold_prefix="")
bullet("再生能源 Ir：太陽能，輸出受日前預測曲線上限約束。")
bullet("儲能裝置 Ib：電池，可充可放但不可同時，受 SOC 上下限約束。")
spacer(2)
para("三類用電需求（任務）：", bold=True)
bullet("週期任務 Jp：固定週期釋放的硬截止任務（例如工廠定時製程）。", bold_prefix="")
bullet("零星任務 Js：硬截止、臨時出現，必須通過「接收測試」才接受（例如緊急冷卻）。")
bullet("非週期任務 Ja：軟截止、可等待，逾時只記為 miss 與 tardiness（例如非緊急充電）。")
spacer(2)
para("系統需同時追求三個目標（彼此衝突）：", bold=True)
numbered("最小化非週期任務逾時數（每漏一個罰 α = 10000 美元）。")
numbered("最小化傳統機組發電成本。")
numbered("最大化電力市場售電收益。")
para("整體目標函數（最小化形式）：", bold=True)
code("min  F = α·f1 + f2 + f3\n"
     "      f1 = Σ_{j∈Ja} Miss_j                         （非週期逾時數，α=10000）\n"
     "      f2 = Σ_{i∈Ig} Σ_t (coston_i·min(1,P) + costup_i·P)  （機組成本）\n"
     "      f3 = − Σ_t (λ_t · Sell_t)                    （售電收益取負號）")

# ---- 2. 輸入 ----
h2("2. 系統輸入與資產設定")
para("輸入檔 input/processor_settings.json 與 input/price_72hr.json 定義固定的資產與價格，"
     "排程時程為 72 小時、每格 Δt = 1 小時。本次使用的資產如下。")
h3("2.1 傳統機組（2 部）")
table(
    ["機組", "出力下/上限(MW)", "爬升率 RU/RD", "最小開/關機(h)", "固定成本($/h)", "變動成本($/MWh)"],
    [["thermal_1", "15 / 80", "15 / 15", "3 / 2", "1200", "42"],
     ["thermal_2", "10 / 45", "20 / 20", "2 / 2", "600", "70"]],
    font=9.0,
)
para("thermal_1 變動成本低（基載）、thermal_2 啟動便宜但每 MWh 較貴（尖峰備用）。兩部初始皆為關機。",
     size=9.5, italic=True)
h3("2.2 再生能源（太陽能 2 部）與儲能（電池 2 部）")
table(
    ["設備", "型別", "關鍵參數"],
    [["pv_1", "太陽能", "額定 60 MW，依日前預測比例 renewforecast 出力"],
     ["pv_2", "太陽能", "額定 80 MW，預測曲線每日三段日照尖峰（約 8–13、31–43、55–66 時）"],
     ["battery_1", "儲能", "SOC 20–100 MWh，充/放 20 MW，初始 45 MWh"],
     ["battery_2", "儲能", "SOC 10–60 MWh，充/放 15 MW，初始 25 MWh"]],
    widths=[1.0, 0.9, 4.2], font=9.0,
)
para("每個電池對應一個「充電需求」job（battery_1_chg / battery_2_chg），充電只能由機組或再生能源供應、"
     "不能由其他電池放電供應。", size=9.5, italic=True)
h3("2.3 市場價格")
para("price_72hr.json 提供每小時售電價 λ_t。本資料尖峰落在第 70 時（113）、第 22 時（108）、"
     "第 56 時（105）等，谷底為第 12 時（0）。排程因此傾向把多餘電力留到高價時段售出。")

# ---- 3. 架構 ----
h2("3. 系統架構與管線（Pipeline）")
para("整個系統由 src/main.py 串接四個階段；Level 2 另以 run_level2() 串成「生成 → 靜態 L1 → 動態 L2 → 比較」。")
code(
    "Phase 1  任務生成   generator/          → output/task_set.json\n"
    "Phase 2  日前排程   rt_scheduler/       → output/schedule_result.json  (PuLP MILP)\n"
    "Phase 3  接收測試   acceptance_tester   → 在排程上標註接受/拒絕 (整合於 Phase 2 之後)\n"
    "Phase 4  效能評估   evaluator/          → output/evaluation_results.json\n"
    "Level 2  動態排程   advanced_scheduler  → output/*_dynamic.json"
)
para("設計上採物件導向與單一職責：生成器把計算交給 FrameSizeCalculator、驗證交給 TaskSetValidator；"
     "排程器把建模交給 VppMilpFormulator、求解結果解析交給 SchedulerResultExtractor。所有資料模型"
     "以 Pydantic 定義，I/O 集中在 JsonIO。")

# ---- 4. Phase 1 ----
h2("4. Phase 1：週期任務集生成")
para("TaskSetGenerator 以「生成 → 計算 frame size → 驗證」的迴圈，隨機產生一組同時滿足所有作業限制的"
     "週期任務；不合格就重抽，直到通過 TaskSetValidator。產生策略刻意保證難約束的條件：")
bullet("先生成「deadline = execution」的任務（保證 1-6：≥20% 任務 d=e）。", bold_prefix="")
bullet("再生成 2 個 e=2 的非搶占任務（保證 1-7：≥2 個 e≠1 的非搶占任務）。")
bullet("其餘任務隨機補足，期間維持週期種類 ≥ 3 種。")
para("Frame size f 的選擇（FrameSizeCalculator）需同時滿足三條件，取最小可行值：")
code("(1) f ≥ max(e_j)        (2) H mod f = 0 (H=72)\n"
     "(3) 2f − gcd(f, p_j) ≤ d_j  對所有任務成立")
h3("4.1 本次產生的週期任務集（9 個任務，frame size = 4）")
table(
    ["ID", "r 釋放", "p 週期", "e 執行", "d 相對截止", "w 能量(MWh)", "搶占", "展開 job 數"],
    [["p1", 23, 24, 4, 4, 17, "是", 2],
     ["p2", 14, 15, 2, 10, 18, "否", 4],
     ["p3", 3, 20, 1, 18, 18, "否", 3],
     ["p4", 19, 20, 4, 4, 14, "否", 3],
     ["p5", 9, 16, 2, 12, 14, "否", 4],
     ["p6", 5, 19, 2, 15, 16, "是", 3],
     ["p7", 14, 17, 3, 11, 6, "是", 3],
     ["p8", 2, 8, 1, 5, 16, "否", 9],
     ["p9", 4, 21, 3, 19, 15, "是", 3]],
    font=9.0,
)
rich([("合計展開為 ", False), ("34 個週期 job", True), ("（> 30 ✔）；密度 D_W = Σ(e/p) = ", False),
      ("1.22 ≥ 0.7 ✔", True), ("；週期種類 8 種（≥3 ✔）；d=e 的有 p1、p4（2/9 ≈ 22% ≥ 20% ✔）；"
       "非搶占且 e≠1 的有 p2、p4、p5（≥2 ✔）。frame size f=4 通過三條件 ✔。", False)])
para("展開規則：第 k 個 job 的 release = r + k·p，deadline = release + d − 1；只保留 deadline ≤ 72 的 job。",
     size=9.5, italic=True)
para("此外 task_set.json 也含 Demo 提供的 6 個零星任務（s1–s6）與 10 個非週期任務（a1–a10），"
     "供 Phase 3 接收測試使用。")

pagebreak()

# ---- 5. Phase 2 ----
h2("5. Phase 2：日前 MILP 排程")
para("VppMilpFormulator 把 72 小時的排程問題建成一個混合整數線性規劃（MILP），以 PuLP + CBC 求解，"
     "一次解出整個 horizon 的最佳排程。")
h3("5.1 決策變數")
table(
    ["變數", "意義"],
    [["P[i][t]", "設備 i 在 t 的總出力（MWh）"],
     ["k[j][i][t]", "設備 i 在 t 供給需求 j 的能量；充電需求則為供給對應電池的充電量"],
     ["u, start, stop[i][t]", "機組開關機 / 啟動 / 停機之二元變數（線性化 min(1,P)）"],
     ["charge_b, discharge_b[i][t]", "電池充 / 放電二元旗標（互斥，線性化 C19）"],
     ["SOC[i][t]", "電池在 t 末的儲能狀態"],
     ["Sell[t]", "t 時售電量"],
     ["x[j][t]", "需求 j 在 t 是否執行之二元變數（線性化 min(1, Σk)）"]],
    widths=[1.6, 4.4], font=9.0,
)
para("作業數學式中的 min(1, ·) 為非線性，實作以二元變數 x（需求是否啟動）與 u（機組是否開機）線性化──"
     "這是把規格轉成可解 MILP 的關鍵，且不新增超出規格語意的決策。", size=9.5, italic=True)

h3("5.2 23 條限制式（依實作分組）")
table(
    ["群組", "編號", "內容（實作對應）"],
    [["任務", "C1–C3, C5",
      "C1 啟動時收到的總能量 = w；C2 釋放前不可執行；C3 在 [r, r+d−1] 內累積執行恰 e 格；"
      "C5 非搶占任務須連續執行（以 0→1 上升沿 ≤1 強制單一區塊）"],
     ["非週期", "C4", "Miss_j 定義：未逾時 = 完成 e 格；逾時則記 Miss=1"],
     ["機組", "C6–C12",
      "C6 出力上下限；C7 爬升率；C8 下限 ≤ 單格爬升；C9/C10 最小開/關機時間；C11/C12 初始開/關機延續"],
     ["再生能源", "C13", "P ≤ capacity · forecast · Δt"],
     ["儲能", "C14–C19, C21",
      "C14 放電上限；C15 充電上限；C16 SOC 平衡；C17 SOC 上下限；C18 放電 ≤ 可用儲能；"
      "C19 不可同時充放；C21 充電不可由非發電設備供應"],
     ["售電", "C22", "Sell_t ≥ 0"],
     ["供需平衡", "C20, C23",
      "C20 各設備供給總和 ≤ 其出力；C23 每小時：總發電 = 需求 + 充電 + 售電"]],
    widths=[0.8, 1.0, 4.2], font=8.5,
)
para("CLAUDE.md 與規格所稱「23 條限制式」即上表所有條目（含同編號的上/下限與爬升上/下兩式）。"
     "C16 與 C18 在 Level 2 會被「真實儲能」放寬版取代（見第 8 節），但在 Level 1 預設參數下與原式等價。",
     size=9.5, italic=True)

h3("5.3 目標函數與求解")
para("目標即第 1 節的 F。Level 1 的 MILP 只將週期 job 與充電 job 納入模型；零星／非週期任務由 Phase 3 "
     "以保留量處理，因此靜態 MILP 的 f1 在求解時為 0，非週期逾時的罰則在 Phase 4 評估時才計入 objective_value。")
para("以 PuLP 呼叫 CBC 求得最佳解後，SchedulerResultExtractor 解析各變數、四捨五入、捨去近零值，"
     "輸出 schedule_result.json（每個 t 含 P / k / sell / soc 及接收測試標註欄位）。")

# ---- 6. Phase 3 ----
h2("6. Phase 3：接收測試（Sporadic／Aperiodic）")
para("AcceptanceTester 在 MILP 解出後執行（已整合進 RTScheduler.run()）。它利用 MILP 留下的「保留量"
     "（reserve）」──即原本要賣到市場的剩餘電力 sell[t]，並可拆解為各設備的 spare = P[i][t] − Σ_j k[j][i][t]"
     "（由 C23 知其總和 = sell[t]）。")
para("接受一個 job 時，把它每格的需求 w 從設備 spare 轉入 k[job][device]，並把 sell[t] 同額減少──"
     "總發電量不變，C23（平衡）與 C20（設備上限）仍成立。決策規則：", )
bullet("零星任務（硬截止）：只有當 [r, r+d−1] 內能排入 e 格、每格 reserve ≥ w 時才接受，否則拒絕。",
       bold_prefix="")
bullet("非週期任務（軟截止）：排到能在 horizon 末 H 前完成；若完成時刻晚於軟截止則另記為 miss。")
bullet("搶占 = 1 可用不連續時槽；搶占 = 0 須單一連續視窗。")
para("每個決策（接受/拒絕、排入的時槽、完成時刻、理由）都寫入 acceptance_test_log.json，直接支援評分 4-1（方法）"
     "與 4-2（決策合理性）。", size=9.5, italic=True)
h3("6.1 本次靜態接收測試結果")
para("零星 s1–s6：全部接受（value rate = 1.0）。例如 s1 在 [2,9] 內排入第 2、4 時；"
     "s3 在 [28,33] 內排入第 28、29 時。非週期 a1–a10：全部排入，但 a2（完成於 21 > 軟截止 17）與 "
     "a5（完成於 46 > 軟截止 42）逾時，故靜態軟逾時率 = 2/10 = 0.2。")

# ---- 7. Phase 4 ----
h2("7. Phase 4：效能評估指標")
para("Evaluator 讀入排程結果與輸入檔，計算所有規格要求的指標，輸出 evaluation_results.json：")
table(
    ["指標", "定義"],
    [["hard_deadline_miss_rate", "（逾時的週期 + 零星 job 數）/ 總硬截止 job 數"],
     ["soft_deadline_miss_rate", "逾時的非週期 job 數 / 總非週期 job 數"],
     ["average / max_tardiness", "T_j = max(0, C_j − d_j) 的平均與最大"],
     ["average / max_response_time", "R_j = C_j − r_j 的平均與最大"],
     ["completion_time_jitter", "同一週期任務各 instance 完成時刻的峰對峰值，再取平均"],
     ["sporadic_value_rate", "在硬截止前完成的零星執行格數 / 零星總執行格數"],
     ["generator_cost / market_revenue", "f2 機組成本 / Σ(λ_t·Sell_t) 售電收益"],
     ["objective_value", "F = α·(軟逾時數) + 機組成本 − 售電收益"]],
    widths=[2.2, 3.8], font=9.0,
)

pagebreak()

# ---- 8. Level 2-A 放寬假設 ----
h2("8. Level 2-A：放寬假設建模（R1–R10）")
rich([("我們放寬三個 Level 1 假設──假設 11（再生能源預測必準）、假設 12（理想儲能）、假設 5（無工作優先序）──"
       "並嚴格遵守規格規定：", False), ("「可修改符號、可加參數，但不得新增決策變數」", True),
      ("。下列每條放寬限制都只用 Level 1 既有變數 P / k / SOC / Sell / x 表達。", False)])
para("所有放寬參數集中在 runtime_config.json；參數全取預設（η=1, σ=0, β=0…）時，模型與 Level 1 完全相同。",
     size=9.5, italic=True)
h3("8.1 放寬參數表（本次採用值）")
table(
    ["參數", "符號", "值", "對應假設", "意義"],
    [["charge_efficiency", "η_c", "0.95", "12", "充電能量僅 η_c 進入電芯"],
     ["discharge_efficiency", "η_d", "0.95", "12", "放出 P 需從電芯抽 P/η_d"],
     ["self_discharge_rate", "σ", "0.002", "12", "每格自放電比例"],
     ["cycle_limit", "N", "3.0", "12", "全期最大等效循環數"],
     ["soc_power_floor", "φ", "0.5", "12", "SOC 極端時的最低功率比例"],
     ["aging_cost", "c_age", "2.0", "12", "每放電 1 MWh 的老化成本($)"],
     ["renewable_uncertainty_margin", "β", "0.15", "11", "預測折減保留邊際"],
     ["precedence", "—", "自動", "5", "有序工作對 (a,b)"]],
    font=9.0,
)
h3("8.2 放寬限制式 R1–R10")
para("假設 11 — 再生能源不確定性", bold=True, color=BLUE)
bullet("R1 / C13′：P[i][t] ≤ capacity · forecast · (1−β) · Δt。日前計畫不靠滿預測，保留邊際以防高估。",
       bold_prefix="")
para("假設 12 — 真實儲能（就地改寫 C16、C18，預設參數即還原）", bold=True, color=BLUE)
bullet("R2/R3/R4（C16′ SOC 平衡）：SOC[t] = (1−σ)·SOC[t−1] + η_c·charge_in − (1/η_d)·P[t]。",
       bold_prefix="")
bullet("R5（C18′）：(1/η_d)·P[t] ≤ (1−σ)·SOC[t−1] − soc_min。")
bullet("R6（循環/吞吐上限）：Σ_t P[i][t] ≤ N·(soc_max − soc_min)。")
bullet("R7（SOC 相依放電功率）：接近滿電時可達額定，接近空電時降到 φ·額定，線性內插。")
bullet("R8（SOC 相依充電功率）：對稱地，接近滿電時充電功率降到 φ·額定。")
bullet("R9（老化成本，進目標）：目標加上 Σ c_age·P[storage]，讓最佳化在「循環電池」與「燒燃料」間權衡。")
para("假設 5 — 工作優先序", bold=True, color=BLUE)
bullet("R10：有序對 (a,b)，在 b 的視窗內 e_a·x[b][t] ≤ Σ_{s<t} x[a][s]，即 a 全部做完前 b 不得啟動；"
       "重用 x，無新變數。未設定時程式自動挑一組可行的非平凡對（本次為 p8_0 → p3_0）。", bold_prefix="")
para("保留策略（支援 6/7 與動態法）", bold=True, color=BLUE)
bullet("R-reserve（保留下限）：Sell[t] ≥ R_t。以對 Sell 設下限的方式，強制計畫保留可轉移的剩餘電力，"
       "供接收測試使用；重用 Sell，無新變數。", bold_prefix="")
rich([("合計 ", False), ("10 條放寬限制（R1–R10）", True),
      (" 外加保留下限，達到評分 3 的上限 10 分，且皆未新增決策變數。", False)])

# ---- 9. Level 2-B 動態排程 ----
h2("9. Level 2-B：進階動態排程（滾動視窗再最佳化）")
para("方法：Receding-horizon（滾動視窗）再最佳化。", bold=True)
para("Level 1 在 t=0 一次解完整個 [1,72]；Level 2 的 AdvancedScheduler 則沿時間前進，在觸發點對「剩餘 horizon」"
     "重新最佳化，藉此因應靜態計畫看不到的動態資訊。")
h3("9.1 核心機制")
bullet("狀態承接（pin_prefix）：每次觸發時，已執行的 [1, t0) 被釘住為已承諾值（連續變數 P/k/SOC/Sell 加上"
       "由其推導的前綴二元變數 u/charge_b/discharge_b/x），只重解 [t0, 72]。釘住前綴讓 solver 的 presolve "
       "把它消去，使每次再解都便宜，且機組開關機時長、爬升位置、SOC 都正確跨界承接。", bold_prefix="")
bullet("觸發時機：(a) 週期節奏 reopt_interval = 24 格；(b) 每個零星任務釋放（硬任務需即時審核）；"
       "(c) 最多 3 個「實際 PV 與預測偏差最大」的時點。非週期任務於下一個節奏邊界揭露。")
bullet("再生能源實現：以種子化隨機模型（noise std = 0.2, seed = 42）在預測周圍生成「實際」PV。"
       "已承諾區塊用實際值、未見的尾段用折減預測（R1）──對已承諾區塊精確、對未來保守。")
bullet("線上接收 = 保留可行性：任務在釋放時揭露，能否接受取決於再最佳化「能否在其執行視窗內保留足夠"
       "可轉移剩餘（Sell ≥ R_t）」。硬任務不可行就拒絕；軟任務盡力保留。被接受者隨區塊凍結被路由進排程，"
       "故接受的硬任務必在截止前完成。")
para("每次再解使用 MIP gap（gap_rel = 0.05）與時間上限（time_limit = 20s），確保線上再最佳化可解。",
     size=9.5, italic=True)
h3("9.2 目標衝突（為何無法同時最佳化三目標）")
bullet("提高零星/非週期接收（降 miss）⇒ 計畫須保留更多剩餘 ⇒ 更多機組承諾 ⇒ 機組成本 f2 上升、"
       "或售電收益下降（剩餘被留著而非在尖峰賣出）。", bold_prefix="")
bullet("再生能源不確定（R1 折減 + 實際短缺）⇒ 能量從 PV 移到火力 ⇒ f2 上升。")
bullet("儲能真實化（效率、自放電、循環、老化）⇒ 電池成為更貴、更有損的緩衝 ⇒ 套利收益變少。")
para("因此動態法是「以較高機組成本換取更少逾時與更佳即時回應」，下一節以實際數據驗證。")

pagebreak()

# ---- 10. 結果 ----
h2("10. 實際執行結果與靜態 vs 動態比較")
h3("10.1 動態排程的 9 輪再最佳化（共 9 次 MILP 解，約 15 秒）")
table(
    ["觸發 t0", "承諾區間", "本輪揭露的任務", "結果"],
    [["1", "[1,1]", "（無）", "建立初始計畫"],
     ["2", "[2,13]", "s1, a1, a2", "全部接受並路由"],
     ["14", "[14,24]", "s2, a3", "全部接受"],
     ["25", "[25,27]", "a4", "接受"],
     ["28", "[28,39]", "s3, a5", "全部接受"],
     ["40", "[40,48]", "s4, a6", "全部接受"],
     ["49", "[49,52]", "a7", "接受（保留下限 7.0）"],
     ["53", "[53,65]", "s5, a8, a9", "全部接受"],
     ["66", "[66,72]", "s6, a10", "全部接受"]],
    widths=[0.8, 1.1, 2.0, 2.1], font=9.0,
)
para("9 輪皆無拒絕、無再排程失敗；6 個零星任務全部於硬截止前完成、10 個非週期任務全部於軟截止前完成。"
     "逐格驗證 C23 平衡 0 違反、SOC 維持在 [soc_min, soc_max] 0 違反、放寬儲能/機組/再生能源限制每輪皆成立。")

h3("10.2 靜態（L1）vs 動態（L2）效能比較")
table(
    ["指標", "靜態 L1", "動態 L2", "差異 Δ"],
    [["objective_value", "−72475.6", "−88505.4", "−16029.8（更佳）"],
     ["generator_cost", "296580", "305700", "+9120"],
     ["market_revenue", "389055.6", "394205.4", "+5149.8"],
     ["hard_deadline_miss_rate", "0.0", "0.0", "—"],
     ["soft_deadline_miss_rate", "0.2", "0.0", "−0.2"],
     ["average_tardiness", "0.16", "0.0", "−0.16"],
     ["max_tardiness", "4", "0", "−4"],
     ["average_response_time", "4.9", "4.36", "−0.54"],
     ["sporadic_value_rate", "1.0", "1.0", "—"]],
    font=9.0,
)
para("解讀：", bold=True)
bullet("靜態日前計畫雖一次看到所有任務，但只承諾一份排程，漏掉 2 個非週期任務（a2、a5，軟逾時率 0.2）。"
       "動態法隨任務到達再最佳化並為其保留容量，10 個全部準時完成（軟逾時率 0.0），平均回應時間由 4.9 降到 4.36。",
       bold_prefix="")
bullet("消除 2 個軟逾時 = 省下 2 × α = 20000 罰則，主導了 objective 的改善（−72476 → −88505）。")
bullet("代價清晰可見：保留容量並補償實際再生能源短缺使機組成本上升（+9120）；較大的承諾發電同時讓售電收益"
       "略增（+5150）。扣掉逾時罰則後，動態法在總目標上明顯更佳，印證 9.2 的權衡分析。")

# ---- 11. 程式結構 ----
h2("11. 程式碼結構與執行方式")
code(
    "src/\n"
    "├── main.py              管線進入點（run_level2 = L1 + L2 + 比較）\n"
    "├── config.py            路徑與常數\n"
    "├── validator.py         自我檢核（純標準庫）\n"
    "├── advanced_scheduler.py  Level 2 滾動視窗動態排程\n"
    "├── generator/           Phase 1 任務生成\n"
    "├── rt_scheduler/        Phase 2+3：formulator / acceptance_tester /\n"
    "│                         relaxation / expander / extractor\n"
    "├── evaluator/           Phase 4 評估\n"
    "└── model/               Pydantic 資料模型\n"
    "input/  processor_settings.json, price_72hr.json\n"
    "output/ task_set.json, schedule_result*.json, evaluation_results*.json,\n"
    "        acceptance_test_log*.json, dynamic_run_log.json\n"
    "runtime_config.json      Level 2 放寬 + 動態節奏參數"
)
para("執行指令：", bold=True)
code(
    "# 環境\n"
    "pip install pulp pydantic        # 或 poetry install\n\n"
    "# Level 1 完整管線（生成 → 排程 → 評估）\n"
    "python -m src.main\n\n"
    "# Level 2 比較管線（生成 → 靜態 → 動態 → 印出比較）\n"
    "python -c \"from src.main import run_level2; run_level2()\"\n\n"
    "# 只跑動態排程 / 自我檢核\n"
    "python -m src.advanced_scheduler\n"
    "python3 -m src.validator            # Level 1\n"
    "python3 -m src.validator --level 2  # Level 2"
)

# ---- 12. 自我檢核 ----
h2("12. 對應評分表的自我檢核")
table(
    ["評分項", "配分", "本作業狀態"],
    [["1 週期任務集（1-1～1-8）", "17", "全數通過（見 4.1）"],
     ["2 模型限制（C1～C23，含放寬版）", "27", "全數實作並通過自我檢核"],
     ["3 放寬假設（R1–R10）", "10", "10 條皆由動態排程實際滿足"],
     ["4 接收測試（含 4-3 value rate）", "11", "零星 value rate = 1.0；附逐 job 理由"],
     ["5 排程結果", "8", "JSON 格式相容；週期 job 全準時"],
     ["6 評估指標", "7", "全部指標皆輸出"],
     ["7 保留策略分析", "10", "報告分析 + 保留下限 R-reserve"],
     ["8 動態法（8-1/8-2/8-3）", "10", "設計 + 正確性 + 靜/動比較皆具備"]],
    widths=[2.6, 0.7, 2.7], font=9.0,
)
para("自我檢核工具 python3 -m src.validator --level 2 會用放寬後的儲能/再生能源限制重驗每一項，"
     "並逐條確認動態排程確實滿足 R1–R10（例如 SOC 是否遵循效率/自放電平衡、再生能源是否守住實際上限、"
     "優先序是否成立），亦即驗證「實作」而非僅「描述」。", size=9.5, italic=True)

# =========================================================================== 第二部分 講稿
pagebreak()
h1("第二部分　報告講稿（明日對老師報告）")
para("建議總時長 10–12 分鐘。以下每段對應一張投影片，粗體為口語講稿，項目為提醒重點。", italic=True)

def script_block(slide_title, minutes, speech, points=None):
    h2(slide_title)
    para(f"（建議 {minutes}）", italic=True, size=9.5, color=BLUE)
    rich([("講稿：", True), (speech, False)])
    if points:
        for pt in points:
            bullet(pt)
    spacer(2)

script_block(
    "投影片 1：開場與題目定位", "約 1 分鐘",
    "老師好，我們這組做的是「虛擬電廠的即時排程系統」。核心想法是：把電廠裡的傳統機組、太陽能、"
    "電池當成即時系統的『處理器』，把各種有時間限制的用電需求當成『任務』，於是供電調度就變成一個"
    "即時排程問題。我們要在 72 小時、每小時一格的時間軸上，安排這些處理器供電，"
    "同時滿足三個彼此衝突的目標：少漏非週期任務、降低機組成本、提高售電收益。",
    ["先讓老師建立『處理器 = 發電設備、任務 = 用電需求』的對應。",
     "強調三個目標互相衝突，這是後面 Level 2 分析的伏筆。"],
)

script_block(
    "投影片 2：系統架構與四階段管線", "約 1 分鐘",
    "整個系統分四個階段：Phase 1 生成週期任務集；Phase 2 用混合整數線性規劃解出日前排程；"
    "Phase 3 在排程留下的保留量上做接收測試，處理零星與非週期任務；Phase 4 計算所有效能指標。"
    "Level 2 我們在這個基礎上放寬三個假設，並加上一個會隨資訊滾動再最佳化的動態排程器。",
    ["指著架構圖由左到右講一次資料流。",
     "點出 Phase 3 已整合進排程器之後自動執行。"],
)

script_block(
    "投影片 3：Phase 1 週期任務生成", "約 1 分鐘",
    "週期任務是用『生成—驗證』迴圈隨機產生的，會反覆重抽直到同時滿足作業所有限制。"
    "我們刻意先生成 d=e 的任務和非搶占任務來保證最難滿足的條件。這次產生 9 個週期任務，frame size 是 4，"
    "展開後是 34 個 job，超過 30 個；密度 1.22 遠高於 0.7 的下限；frame size 也通過 2f − gcd(f,p) ≤ d 這條。",
    ["報出關鍵數字：9 個任務、frame=4、34 個 job、密度 1.22。",
     "若老師問為何是 frame size，說明它保證每個 job 有完整一格可被排進去。"],
)

script_block(
    "投影片 4：Phase 2 MILP 模型──變數與目標", "約 1.5 分鐘",
    "日前排程是一個 MILP，我們用 PuLP 加 CBC 一次解出整個 72 小時。決策變數有設備出力 P、"
    "能量分配 k、機組開關機、電池充放電與 SOC、售電量 Sell，還有需求是否執行的二元變數 x。"
    "目標函數是 α 乘上非週期逾時數，加機組成本，減售電收益；α 設成 10000，代表每漏一個非週期任務罰一萬元。"
    "規格裡的 min(1, ·) 是非線性，我們用二元變數 x 和 u 把它線性化，這樣才能交給 MILP 解。",
    ["強調『一次解完整個 horizon』是靜態法的特徵，與 Level 2 對照。",
     "解釋 α=10000 把『個數』換算成『金額』才能和成本相加。"],
)

script_block(
    "投影片 5：Phase 2 模型──23 條限制式", "約 1.5 分鐘",
    "限制式共 23 條，分成幾群：任務類保證每個 job 在截止前累積執行恰好 e 格、非搶占要連續；"
    "機組類包含出力上下限、爬升率、最小開關機時間；再生能源受預測上限約束；"
    "儲能類包含充放電上限、SOC 平衡、不可同時充放；最後是每小時的供需平衡 C23──"
    "每一格的總發電必須等於用電需求加充電加售電。這條平衡式是整個模型的骨幹。",
    ["不必逐條念，挑 C23 平衡與 C3 完成性、C5 連續性講透即可。",
     "提一句 C16/C18 在 Level 2 會被放寬版取代。"],
)

script_block(
    "投影片 6：Phase 3 接收測試與保留策略", "約 1.5 分鐘",
    "排程解完後，市場上沒賣掉的剩餘電力就是我們的『保留量』。接收測試就靠這個保留量來接零星和非週期任務。"
    "零星任務是硬截止：只有當它的時間視窗內每一格的保留量都夠它的能量需求、而且能排滿執行格數時，才接受，"
    "否則拒絕。非週期任務是軟截止：盡量排，逾時就記成 miss。接受時我們把需求從設備的剩餘出力轉進去、"
    "同時把售電量等額扣掉，所以總發電不變、平衡式仍然成立。這次 6 個零星任務全部接受，value rate 是 1.0。",
    ["講清楚『轉移剩餘、扣售電、平衡不變』這個關鍵手法。",
     "報出：零星全接受、value rate 1.0；靜態下 a2、a5 兩個非週期逾時。"],
)

script_block(
    "投影片 7：Level 2-A 放寬三個假設", "約 1.5 分鐘",
    "Level 2 我們放寬三個假設：再生能源預測不一定準、儲能不是理想的、以及工作之間有優先序。"
    "最重要的限制是：規格規定可以加參數、改符號，但不能新增決策變數，所以我們所有放寬限制都只用"
    "原本的變數來寫。再生能源我們把預測折減 15% 留邊際；儲能我們加了充放電效率 0.95、自放電、"
    "循環上限、SOC 相依功率和老化成本；優先序則是強制某個 job 做完另一個才能開始。"
    "總共 10 條放寬限制，剛好對到評分的 10 分上限。把參數設回預設，模型就跟 Level 1 完全一樣。",
    ["反覆強調『沒有新增決策變數』──這是評分硬規定。",
     "舉一兩個具體：SOC 平衡加了效率與自放電；老化成本進目標。"],
)

script_block(
    "投影片 8：Level 2-B 動態排程方法", "約 1.5 分鐘",
    "動態排程用的是滾動視窗再最佳化。Level 1 是一次解完全程，動態法則是沿時間往前走，在三種觸發點"
    "重新最佳化剩下的時程：固定每 24 小時一次、每個零星任務一到就審核、還有實際太陽能跟預測差最大的幾個點。"
    "已經執行過的部分會被『釘住』，只重解後面，這樣每次再解都很快，而且機組狀態跟 SOC 都正確接續。"
    "已承諾的區塊用『實際』的太陽能出力、看不到的未來用折減預測，所以對已承諾精確、對未來保守。"
    "任務能不能接受，看再最佳化能不能在它的視窗內保留足夠剩餘，這就是把接收測試做成 MILP 的可行性問題。",
    ["強調『釘住前綴 + 只重解尾段』是讓線上再最佳化可解的關鍵。",
     "用一句話帶出：硬任務接受了就保證準時完成。"],
)

script_block(
    "投影片 9：結果──9 輪再最佳化與正確性", "約 1 分鐘",
    "這次動態排程跑了 9 輪、9 次 MILP，大約 15 秒。任務隨時間陸續到達並被接受："
    "例如第 2 小時揭露 s1、a1、a2，第 28 小時揭露 s3、a5，一路到第 66 小時的 s6、a10。"
    "最後所有 34 個週期 job、6 個零星 job 都在硬截止前完成，10 個非週期 job 也都在軟截止前完成；"
    "硬逾時率 0、軟逾時率 0。每一格的供需平衡和 SOC 範圍我們也都逐格驗證沒有違反。",
    ["報出：9 輪、9 次解、約 15 秒、全部準時。",
     "強調『逐格驗證 0 違反』回應評分 8-2 的正確性要求。"],
)

script_block(
    "投影片 10：結果──靜態 vs 動態比較", "約 1.5 分鐘",
    "最關鍵的比較在這裡。靜態法雖然一次看到所有任務，但只承諾一份計畫，結果漏掉 2 個非週期任務，"
    "軟逾時率 0.2。動態法隨任務到達再最佳化、並主動保留容量，10 個非週期任務全部準時，軟逾時率變成 0，"
    "平均回應時間也從 4.9 降到 4.36。少漏 2 個就省下 2 萬元罰則，這主導了總目標從 −72476 改善到 −88505。"
    "代價是：保留容量加上補償太陽能實際短缺，機組成本多了約 9 千；不過較大的發電也讓售電收益增加約 5 千。"
    "這正好驗證我們前面說的──提高接收率要用機組成本來換，三個目標沒辦法同時最佳化。",
    ["這是全場重點，務必把『省 2 萬罰則 > 多花 9 千成本』的帳算給老師聽。",
     "結論句：動態法在總目標上明顯更好，且權衡關係與理論分析一致。"],
)

script_block(
    "投影片 11：結語與分工／AI 協作", "約 1 分鐘",
    "總結：我們完成了 Level 1 的生成、日前排程、接收測試與評估，也完成 Level 2 的三項假設放寬與動態排程，"
    "並用實際數據證明動態法在不漏硬截止的前提下，把軟逾時降到 0、總目標更好。"
    "程式採物件導向、單一職責設計，所有結果都可用固定種子重現。謝謝老師，以下開放提問。",
    ["若需報告 AI 協作：說明用 AI 輔助建模與程式骨架、自己負責驗證與調參，並交叉檢查每條限制式。",
     "預留時間回答提問。"],
)

h2("預期問答（備用）")
bullet("Q：為什麼用 MILP 而不是啟發式？　A：MILP 能在 23 條硬限制下保證可行且最佳，"
       "且線上再最佳化用 gap 與時間上限即可在秒級內解出。", bold_prefix="")
bullet("Q：動態法的計算成本？　A：本例 9 次解約 15 秒；釘住前綴讓 presolve 消去已承諾段，"
       "再解規模隨時間縮小。")
bullet("Q：保證沒有新增決策變數？　A：所有放寬限制（R1–R10）與保留下限都只用 P/k/SOC/Sell/x，"
       "已用自我檢核逐條驗證。")
bullet("Q：completion_time_jitter 動態略增（40→42.3）為何仍較好？　A：jitter 是同任務各 instance "
       "完成時刻的離散度，動態法為保留容量微調了部分週期 job 的完成時刻，但硬截止全數守住、"
       "且總目標更佳，屬可接受的取捨。")

spacer(6)
para("（報告結束）", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

# =========================================================================== 存檔
out = "報告_VPP-RTS.docx"
doc.save(out)
print(f"saved: {out}")
